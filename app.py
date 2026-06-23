import streamlit as st
import math
import json
import os
import time
import random
import io
import sqlite3
from datetime import date, datetime, timedelta
import pandas as pd
import numpy as np
import plotly.graph_objects as go
import plotly.express as px
from requests import post
from PIL import Image, ImageDraw, ImageFont

# 加载环境变量
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    if os.path.exists('.env'):
        with open('.env', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#') and '=' in line:
                    k, v = line.split('=', 1)
                    os.environ[k.strip()] = v.strip()

DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "")

# ============================================================
# 加载 JSON 配置文件
# ============================================================
with open("questions.json", encoding="utf-8") as f:
    CONFIG = json.load(f)

ELEMENTS = CONFIG["elements"]       # 四元素配置（水/火/风/土）
QUESTIONS = CONFIG["questions"]     # 测试问题列表
COMPOUNDS = CONFIG["compounds"]     # 元素组合配置
SUGGESTIONS = CONFIG["suggestions"]  # 个性化建议
COMBO_SUG = CONFIG["combo_suggestions"]  # 组合建议
# 成就系统已放弃
ENERGY_TIPS = CONFIG["energy_tips"]      # 每日能量小贴士

# 全局常量定义
ELEMENT_ORDER = ["水", "火", "风", "土"]  # 元素顺序
RADIO_OPTS = ["完全不赞同", "不太赞同", "中立", "比较赞同", "完全赞同"]  # 单选选项
RADIO_MAP = {"完全不赞同": 2, "不太赞同": 4, "中立": 6, "比较赞同": 8, "完全赞同": 10}  # 选项分值映射

# ============================================================
# 数据库操作模块
# ============================================================
DB_PATH = "element_history.db"


def init_db():
    """初始化测试记录数据库表"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS test_records (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nickname TEXT,
            timestamp TEXT,
            dominant_element TEXT,
            scores_water REAL,
            scores_fire REAL,
            scores_wind REAL,
            scores_earth REAL,
            raw_answers TEXT
        )
    ''')
    conn.commit()
    conn.close()


init_db()


def save_record(nickname, dominant, scores_vec, answers_json):
    """保存测试记录到数据库"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute(
        '''
        INSERT INTO test_records (nickname, timestamp, dominant_element,
                                  scores_water, scores_fire, scores_wind, scores_earth, raw_answers)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
    ''',
        (nickname,
         datetime.now().isoformat(),
         dominant,
         scores_vec[0],
         scores_vec[1],
         scores_vec[2],
         scores_vec[3],
         answers_json))
    conn.commit()
    conn.close()


def get_history(nickname=None, limit=10):
    """获取历史测试记录"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    if nickname:
        c.execute('''
            SELECT timestamp, dominant_element, scores_water, scores_fire, scores_wind, scores_earth
            FROM test_records WHERE nickname=? ORDER BY id DESC LIMIT ?
        ''', (nickname, limit))
    else:
        c.execute('''
            SELECT timestamp, dominant_element, scores_water, scores_fire, scores_wind, scores_earth
            FROM test_records ORDER BY id DESC LIMIT ?
        ''', (limit,))
    rows = c.fetchall()
    conn.close()
    return rows


def init_community_db():
    """初始化社区帖子数据库表"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS community_posts (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nickname TEXT DEFAULT "匿名旅人",
            dominant_element TEXT,
            scores_water REAL,
            scores_fire REAL,
            scores_wind REAL,
            scores_earth REAL,
            compound_name TEXT,
            message TEXT,
            timestamp TEXT,
            likes INTEGER DEFAULT 0
        )
    ''')
    # 兼容旧数据库结构，尝试添加缺失字段
    try:
        c.execute('ALTER TABLE community_posts ADD COLUMN likes INTEGER DEFAULT 0')
    except sqlite3.OperationalError:
        pass
    try:
        c.execute(
            'ALTER TABLE community_posts ADD COLUMN nickname TEXT DEFAULT "匿名旅人"')
    except sqlite3.OperationalError:
        pass
    # 为旧数据填充随机点赞数
    c.execute(
        'UPDATE community_posts SET likes = ABS(RANDOM() % 46) + 5 WHERE likes = 0')
    conn.commit()
    conn.close()


init_community_db()


def save_post(nickname, dominant, scores_vec, compound_name, message):
    """保存社区帖子"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    init_likes = random.randint(5, 50)
    c.execute('''
        INSERT INTO community_posts (nickname, dominant_element, scores_water, scores_fire,
                                      scores_wind, scores_earth, compound_name, message, timestamp, likes)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
    ''', (nickname, dominant, scores_vec[0], scores_vec[1], scores_vec[2], scores_vec[3],
          compound_name, message, datetime.now().isoformat(), init_likes))
    conn.commit()
    conn.close()


def get_posts(limit=50):
    """获取社区帖子列表"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute('''
        SELECT id, nickname, dominant_element, scores_water, scores_fire, scores_wind,
               scores_earth, compound_name, message, timestamp, likes
        FROM community_posts ORDER BY id DESC LIMIT ?
    ''', (limit,))
    rows = c.fetchall()
    conn.close()
    return rows


def like_post(post_id):
    """点赞社区帖子"""
    conn = sqlite3.connect(DB_PATH)
    c = conn.cursor()
    c.execute(
        'UPDATE community_posts SET likes = likes + 1 WHERE id = ?', (post_id,))
    conn.commit()
    conn.close()


# ============================================================
# 会话状态初始化
# ============================================================
# 初始化核心会话状态变量
for key in ["page", "step", "best_match", "scores", "distances",
            "future", "records", "chat_history", "prev_page",
            "checked_in_today", "report_generated",
            "nickname", "pending_toast"]:
    if key not in st.session_state:
        if key == "chat_history":
            st.session_state[key] = []
        elif key in ("report_generated", "checked_in_today"):
            st.session_state[key] = False
        elif key in ("scores", "distances"):
            st.session_state[key] = {}
        elif key == "step":
            st.session_state[key] = 0
        elif key == "page":
            st.session_state[key] = "test"
        elif key == "prev_page":
            st.session_state[key] = "test"
        elif key == "records":
            st.session_state[key] = None
        else:
            st.session_state[key] = ""

# 额外的会话状态初始化
if "answer" not in st.session_state:
    st.session_state.answer = None
if "liked_posts" not in st.session_state:
    st.session_state.liked_posts = set()

# ============================================================
# 评分引擎模块
# ============================================================


def calc_scores():
    """遍历所有题目，按元素累加得分"""
    totals = {e: 0 for e in ELEMENT_ORDER}
    all_scores, question_texts = [], []

    for q in QUESTIONS:
        key = f"q{q['id']:02d}"
        val = st.session_state.get(key)

        if val is None:
            s = 0
        elif q["type"] == "radio":
            s = RADIO_MAP.get(val, 0)
        else:
            s = int(val)

        totals[q["element"]] += s
        all_scores.append(s)
        question_texts.append(q["text"])

    return [totals[e] for e in ELEMENT_ORDER], (all_scores, question_texts)


def compute_match(user_vec):
    """计算用户与各元素理想向量的欧氏距离，返回最近元素"""
    profiles = {}
    for idx, elem in enumerate(ELEMENT_ORDER):
        info = ELEMENTS[elem]
        dominant = info["profile_weight"]
        non_dominant = int(info["max_score"] * 0.2)  # 14
        profiles[f"{info['emoji']} {elem}"] = [
            dominant if i == idx else non_dominant for i in range(4)
        ]

    distances, best, best_d = {}, "", float('inf')
    for name, p in profiles.items():
        d = math.sqrt(sum((user_vec[i] - p[i]) ** 2 for i in range(4)))
        distances[name] = d
        if d < best_d:
            best, best_d = name, d
    return best, distances


def get_compound(e1, e2, e3=None):
    """获取元素组合信息"""
    # 先尝试双元素组合
    candidates = [c for c in COMPOUNDS if set(c["combo"]) == {e1, e2}]
    if candidates:
        return candidates[0]
    # 如果有第三个元素，尝试三元素组合
    if e3:
        candidates = [c for c in COMPOUNDS if set(c["combo"]) == {e1, e2, e3}]
        if candidates:
            return candidates[0]
    return None


def get_suggestions(scores_vec):
    """获取个性化成长建议"""
    results = {}
    for i, elem in enumerate(ELEMENT_ORDER):
        # 计算该元素得分百分比
        pct = scores_vec[i] / ELEMENTS[elem]["max_score"]
        # 根据高低分返回不同建议
        tag = "high" if pct >= 0.5 else "low"
        results[elem] = SUGGESTIONS[elem][tag]
    return results


def get_energy_tip(elem):
    """获取今日能量小贴士（每日刷新）"""
    today = date.today().isoformat()
    # 如果日期变更，重新获取小贴士
    if st.session_state.get("_tip_date") != today:
        tip = random.choice(ENERGY_TIPS[elem])
        st.session_state["_tip_date"] = today
        st.session_state["_tip"] = tip
    return st.session_state.get("_tip", ENERGY_TIPS[elem][0])

# ============================================================
# 生成结果分享图片
# ============================================================


def generate_result_image(nickname, element, scores_vec, compound_name=None):
    """生成测试结果分享图片"""
    width, height = 800, 600
    img = Image.new('RGBA', (width, height), color=(20, 18, 38, 255))
    draw = ImageDraw.Draw(img)

    # 渐变背景
    for y in range(height):
        r = int(20 + (40 - 20) * y / height)
        g = int(18 + (20 - 18) * y / height)
        b = int(38 + (80 - 38) * y / height)
        draw.line([(0, y), (width, y)], fill=(r, g, b))

    # 装饰性光晕
    elem_short = element.split()[-1] if " " in element else element
    elem_color = ELEMENTS[elem_short]["color"]
    ec = tuple(int(elem_color.lstrip('#')[j:j + 2], 16) for j in (0, 2, 4))
    for r in range(200, 0, -8):
        alpha = max(0, 8 - (200 - r) // 30)
        draw.ellipse([width // 2 - r, 60 - r, width // 2 + r, 60 + r],
                     fill=(ec[0], ec[1], ec[2], alpha))

    # 星星装饰
    for _ in range(80):
        cx = random.randint(0, width)
        cy = random.randint(0, height)
        s = random.choice([2, 3, 4])
        a = random.randint(15, 50)
        draw.ellipse([cx, cy, cx + s, cy + s], fill=(255, 255, 255, a))

    # 加载字体
    try:
        font_title = ImageFont.truetype("simhei.ttf", 36)
        font_subtitle = ImageFont.truetype("simhei.ttf", 24)
        font_text = ImageFont.truetype("simhei.ttf", 20)
        font_small = ImageFont.truetype("simhei.ttf", 16)
        font_badge = ImageFont.truetype("simhei.ttf", 14)
    except BaseException:
        font_title = ImageFont.load_default()
        font_subtitle = ImageFont.load_default()
        font_text = ImageFont.load_default()
        font_small = ImageFont.load_default()
        font_badge = ImageFont.load_default()

    # 顶部装饰渐变线
    for i in range(width):
        c = int(255 - 200 * abs(i - width / 2) / (width / 2))
        draw.point((i, 0), fill=(c, 215, 0))
        if 5 < i < width - 5:
            draw.point((i, 1), fill=(c // 2, 215 // 2, 0))

    # 标题
    draw.text((width // 2, 42), f"{nickname} 的元素人格",
              fill=(255, 215, 0), anchor="mt", font=font_title)

    # 主导元素发光卡片
    ec_dark = (ec[0] // 6, ec[1] // 6, ec[2] // 6)
    draw.rounded_rectangle([(width // 2 - 160,
                             82),
                            (width // 2 + 160,
                             140)],
                           radius=20,
                           fill=ec_dark,
                           outline=ec,
                           width=2)
    draw.text((width // 2, 111), f"✨ 主导元素：{element}",
              fill=(255, 255, 255), anchor="mt", font=font_subtitle)

    # 各元素得分条（美化版）
    bar_x, bar_y, bar_w, bar_h = 130, 180, 540, 30
    for i, elem in enumerate(ELEMENT_ORDER):
        info = ELEMENTS[elem]
        pct = int(scores_vec[i] / info["max_score"] * 100)
        y = bar_y + i * 52
        color_hex = info["color"].lstrip('#')
        c_rgb = tuple(int(color_hex[j:j + 2], 16) for j in (0, 2, 4))

        # 标签（左侧）
        draw.text((30,
                   y + 5),
                  f"{info['emoji']} {elem}",
                  fill=c_rgb,
                  font=font_text)

        # 背景条（圆角+阴影效果）
        draw.rounded_rectangle([(bar_x, y +
                                 4), (bar_x +
                                      bar_w +
                                      4, y +
                                      bar_h +
                                      4)], radius=16, fill=(c_rgb[0] //
                               8, c_rgb[1] //
                               8, c_rgb[2] //
                               8))
        draw.rounded_rectangle([(bar_x, y), (bar_x +
                                             bar_w, y +
                                             bar_h)], radius=14, fill=(c_rgb[0] //
                                                                       5, c_rgb[1] //
                                                                       5, c_rgb[2] //
                                                                       5))

        # 渐变填充条
        fill_w = int(bar_w * pct / 100)
        if fill_w > 4:
            for x in range(bar_x, bar_x + fill_w):
                ratio = (x - bar_x) / fill_w
                r2 = int(c_rgb[0] * (0.7 + 0.3 * ratio))
                g2 = int(c_rgb[1] * (0.7 + 0.3 * ratio))
                b2 = int(c_rgb[2] * (0.7 + 0.3 * ratio))
                draw.line([(x, y + 2), (x, y + bar_h - 2)],
                          fill=(min(r2, 255), min(g2, 255), min(b2, 255)))
            draw.rounded_rectangle([(bar_x, y), (bar_x + fill_w, y + bar_h)], radius=14,
                                   fill=None, outline=(255, 255, 255, 40), width=1)

        # 百分比徽章
        draw.rounded_rectangle([(bar_x +
                                 bar_w +
                                 10, y -
                                 2), (bar_x +
                                      bar_w +
                                      60, y +
                                      bar_h +
                                      2)], radius=12, fill=(c_rgb[0] //
                               3, c_rgb[1] //
                               3, c_rgb[2] //
                               3))
        draw.text((bar_x + bar_w + 35, y + bar_h // 2), f"{pct}%",
                  fill=c_rgb, anchor="mm", font=font_badge)

    # 组合信息
    if compound_name:
        draw.rounded_rectangle([(width // 2 - 150, 410), (width // 2 + 150, 455)], radius=16,
                               fill=(40, 40, 80), outline=(100, 100, 180), width=1)
        draw.text((width // 2, 432), f"🧬 {compound_name}",
                  fill=(200, 200, 255), anchor="mt", font=font_text)

    # 底部装饰线
    for i in range(width):
        c = int(100 - 80 * abs(i - width / 2) / (width / 2))
        draw.point((i, height - 3), fill=(c, c, c + 40))

    # 底部版权
    draw.text((width // 2, height - 16), "生成于 元素人格测试仪",
              fill=(100, 100, 140), anchor="mb", font=font_small)

    buf = io.BytesIO()
    img.save(buf, format='PNG')
    return buf.getvalue()

# ============================================================
# Word 报告生成
# ============================================================


def generate_word(name, element, vec, distances, suggestions_data):
    """生成 Word 格式的测试报告"""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # 默认样式
    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.font.name = '微软雅黑'
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5

    # 封面标题
    title = doc.add_heading('', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('元素人格测试报告')
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(60, 60, 120)
    run.bold = True

    # 装饰分隔线
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_line = p_line.add_run('━' * 30)
    run_line.font.color.rgb = RGBColor(180, 180, 200)
    run_line.font.size = Pt(8)

    doc.add_paragraph()

    # 基本信息卡片
    elem_short = element.split()[-1] if " " in element else element
    elem_color = ELEMENTS[elem_short]["color"]
    ec = tuple(int(elem_color.lstrip('#')[j:j + 2], 16) for j in (0, 2, 4))

    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ('姓名', name or '匿名'),
        ('报告编号', f'ER-{datetime.now().strftime("%Y%m%d%H%M")}'),
        ('生成日期', datetime.now().strftime('%Y年%m月%d日'))
    ]
    for i, (k, v) in enumerate(info_data):
        c0 = info_table.rows[i].cells[0]
        c1 = info_table.rows[i].cells[1]
        c0.text = k
        c1.text = v
        for p in c0.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
    info_table.columns[0].width = Cm(4)
    info_table.columns[1].width = Cm(8)

    doc.add_paragraph()

    # 主导元素
    p_elem = doc.add_paragraph()
    p_elem.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_label = p_elem.add_run('主导元素：')
    run_label.bold = True
    run_label.font.size = Pt(14)
    run_elem = p_elem.add_run(element)
    run_elem.bold = True
    run_elem.font.size = Pt(18)
    run_elem.font.color.rgb = RGBColor(ec[0], ec[1], ec[2])

    doc.add_paragraph()

    # 元素得分表
    doc.add_heading('📊 元素得分', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    headers = ['元素', '得分', '满分', '占比']
    for i, h in enumerate(headers):
        hdr[i].text = h
        for paragraph in hdr[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    for i, elem in enumerate(ELEMENT_ORDER):
        info = ELEMENTS[elem]
        pct = min(int(vec[i] / info["max_score"] * 100), 100)
        row = table.add_row().cells
        row[0].text = f"{info['emoji']} {elem}"
        row[1].text = f"{vec[i]:.0f}"
        row[2].text = str(info["max_score"])
        row[3].text = f"{pct}%"
        for cell in row:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 成长建议
    doc.add_heading('💡 个性化成长建议', level=1)
    for elem, tips in suggestions_data.items():
        info = ELEMENTS[elem]
        p = doc.add_paragraph()
        run = p.add_run(f"{info['emoji']} {elem}：")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(
            *tuple(int(info['color'].lstrip('#')[j:j + 2], 16) for j in (0, 2, 4)))
        for tip in tips:
            doc.add_paragraph(tip, style='List Bullet')

    # 匹配度
    doc.add_heading('📈 元素匹配度（距离越小越匹配）', level=1)
    table2 = doc.add_table(rows=1, cols=2)
    table2.style = 'Light Shading Accent 1'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    h2 = table2.rows[0].cells
    h2[0].text = '元素类型'
    h2[1].text = '匹配距离'
    for cell in h2:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
    for ename, distance in sorted(distances.items(), key=lambda x: x[1]):
        row = table2.add_row().cells
        row[0].text = ename
        row[1].text = f"{distance:.2f}"
        for cell in row:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(11)

    doc.add_paragraph()

    # 页脚
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('—— 生成于 元素人格测试仪 ——')
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ============================================================
# 助手风格预设
# ============================================================
CHAT_STYLES = {
    "🔥 火之风格": "语气热情、激昂、充满动力。像火焰一样激励行动，强调勇气和突破。",
    "💧 水之风格": "语气温柔、细腻、富有同理心。像流水一样包容，注重情感共鸣和内心成长。",
    "🌬️ 风之风格": "语气灵动、开放、富有思辨。像风一样自由，注重多样化视角和创意启发。",
    "🌍 土之风格": "语气沉稳、务实、条理清晰。像大地一样可靠，强调计划执行和稳步成长。",
}

# ============================================================
# DeepSeek API 调用
# ============================================================


def call_deepseek(messages, web_search=False):
    """调用 DeepSeek API 进行对话"""
    # 检查 API 密钥是否有效
    if not DEEPSEEK_API_KEY or "虚拟" in DEEPSEEK_API_KEY:
        return ("🔮 **当前使用的是虚拟 API 密钥**，无法连接 DeepSeek 服务。\n\n"
                "请将 `.env` 文件中的 `DEEPSEEK_API_KEY` 替换为你的真实密钥。\n"
                "获取地址: https://platform.deepseek.com/api_keys")

    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"}
    data = {"model": "deepseek-chat", "messages": messages, "stream": False}

    if web_search:
        data["enable_search"] = True

    # 重试机制（最多3次）
    for attempt in range(3):
        try:
            resp = post("https://api.deepseek.com/chat/completions",
                        headers=headers, json=data, timeout=30)
            resp.raise_for_status()
            return resp.json()["choices"][0]["message"]["content"]
        except Exception as e:
            if attempt < 2:
                time.sleep(2)
            else:
                return f"❌ API 调用失败（已重试3次）：{e}"


def build_system_prompt(style=list(CHAT_STYLES.keys())[0]):
    """构建系统提示词"""
    scores = st.session_state.get("scores", {})
    best = st.session_state.get("best_match", "")
    style_instruction = CHAT_STYLES.get(style, list(CHAT_STYLES.values())[0])

    if scores:
        scores_str = "; ".join([f"{k}={int(v)}" for k, v in scores.items()])
        return f"你是一位元素人格成长导师。用户的测试结果为：{scores_str}，主导元素：{best}。{style_instruction} 请基于以上数据给出具体建议，无需用户再次说明元素。如用户话题偏离，温和引导回人格成长主题。"
    else:
        return f"你是一位贴心的成长导师。{style_instruction} 与用户进行开放友好的对话，帮助用户探索自我、获得成长启发。"


def dynamic_presets():
    """根据主导元素生成动态快捷问题"""
    best = st.session_state.get("best_match", "")

    # 各元素对应的针对性问题
    base_questions = {
        "水": ["如何增强行动力？", "适合的冥想方式？", "人际关系建议？"],
        "火": ["如何培养耐心？", "团队合作技巧？", "压力管理方法？"],
        "风": ["如何提升专注力？", "职业发展建议？", "如何落实想法？"],
        "土": ["如何变得更灵活？", "创新思维训练？", "如何接受变化？"]
    }

    if best:
        # 提取主导元素名称（去除 emoji）
        elem_name = best.split()[-1]
        return base_questions.get(elem_name, base_questions["水"])

    # 默认问题（未完成测试时）
    return ["如何发现自己的优势？", "怎样平衡工作与生活？", "提升人际关系的方法？"]

# ============================================================
# CSS 样式渲染
# ============================================================


def render_css():
    """渲染自定义 CSS 样式"""
    st.markdown("""
    <style>
    /* 全局背景 */
    .stApp {
        background: linear-gradient(135deg, #0f0c29 0%, #302b63 50%, #24243e 100%);
        background-attachment: fixed;
    }

    /* 标题样式 */
    h1 {
        background: linear-gradient(90deg, #f7971e, #ffd200);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        font-weight: 800 !important;
        text-align: center;
        font-size: 2.6rem !important;
    }

    /* 文本颜色 */
    p, label, .stMarkdown, .stText, h2, h3, h4 {
        color: white !important;
    }

    /* 输入框样式 */
    .stTextInput > div > div > input {
        background: rgba(255,255,255,0.85) !important;
        border-radius: 12px !important;
        color: black !important;
    }

    /* 侧边栏样式 */
    [data-testid="stSidebar"] {
        background: rgba(15,12,41,0.95);
        border-right: 1px solid rgba(255,255,255,0.08);
    }
    [data-testid="stSidebar"] p, [data-testid="stSidebar"] h2,
    [data-testid="stSidebar"] label, [data-testid="stSidebar"] span {
        color: white !important;
    }
    [data-testid="stSidebar"] .stDataFrame {
        background: rgba(255,255,255,0.06) !important;
        border-radius: 10px !important;
    }
    [data-testid="stSidebar"] .stDataFrame td,
    [data-testid="stSidebar"] .stDataFrame th {
        color: white !important;
    }

    /* 按钮样式 */
    .stButton > button, [data-testid="stDownloadButton"] button {
        background: linear-gradient(90deg, #667eea, #764ba2) !important;
        color: white !important;
        border: none !important;
        border-radius: 12px !important;
        font-weight: 600 !important;
        transition: all 0.3s ease !important;
        box-shadow: 0 4px 15px rgba(102,126,234,0.4) !important;
    }
    .stButton > button:hover, [data-testid="stDownloadButton"] button:hover {
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(102,126,234,0.6) !important;
    }

    /* 表单组件容器 */
    [data-testid="stRadio"], .stSelectbox, [data-testid="stSlider"], [data-testid="stNumberInput"] {
        background: rgba(255,255,255,0.07);
        padding: 0.8rem 1rem;
        border-radius: 12px;
        border: 1px solid rgba(255,255,255,0.12);
        margin: 0.4rem 0;
    }

    /* 单选按钮标签 */
    [data-testid="stRadio"] label, [data-testid="stRadio"] label p {
        color: white !important;
    }

    /* 选择框和滑块标签 */
    .stSelectbox label, [data-testid="stSlider"] label, [data-testid="stNumberInput"] label {
        color: white !important;
    }

    /* 数字输入框 */
    [data-testid="stNumberInput"] input {
        background: rgba(255,255,255,0.85) !important;
        color: black !important;
        border-radius: 10px !important;
    }

    /* 数据表格 */
    .stDataFrame {
        background: rgba(255,255,255,0.07) !important;
        border-radius: 12px !important;
    }
    .stDataFrame td, .stDataFrame th {
        color: black !important;
    }

    /* 聊天输入框 */
    [data-testid="stChatInput"] textarea {
        background: rgba(255,255,255,0.85) !important;
        color: black !important;
        border: 1px solid rgba(255,255,255,0.15) !important;
        border-radius: 10px !important;
    }
    [data-testid="stChatInput"] textarea::placeholder {
        color: rgba(0,0,0,0.4) !important;
    }

    /* 聊天消息 */
    [data-testid="stChatMessageContent"] p {
        color: white !important;
    }
    .stChatMessage {
        background: rgba(255,255,255,0.07) !important;
        border-radius: 12px !important;
        padding: 0.3rem 0.8rem !important;
        margin: 0.3rem 0 !important;
    }

    /* 分隔线 */
    hr {
        border-color: rgba(255,255,255,0.08) !important;
        margin: 2rem 0 !important;
    }

    /* 隐藏页脚和菜单 */
    footer, #MainMenu {
        visibility: hidden;
    }

    /* 标签页样式 */
    [data-testid="stTabs"] {
        gap: 0 !important;
    }
    [data-testid="stTabs"] button {
        flex: 1 !important;
        text-align: center !important;
    }
    [data-testid="stTabs"] button p {
        color: white !important;
        font-size: 1.1rem !important;
        letter-spacing: 1.5px !important;
    }

    /* 表格对齐 */
    [data-testid="stTable"] table td, [data-testid="stTable"] table th,
    [data-testid="stDataFrame"] td, [data-testid="stDataFrame"] th {
        text-align: center !important;
    }

    /* 单选按钮布局 */
    div.row-widget.stRadio {
        width: 100% !important;
    }
    div.row-widget.stRadio > div[role="radiogroup"] {
        display: grid !important;
        grid-template-columns: 1fr 1fr !important;
        gap: 0.5rem !important;
        width: 100% !important;
    }
    div.row-widget.stRadio label[data-testid="stRadio"] {
        text-align: center !important;
        justify-content: center !important;
        padding: 0.25rem 0 !important;
    }

    /* 步骤卡片 */
    .step-card {
        background: rgba(255,255,255,0.07);
        backdrop-filter: blur(20px);
        border-radius: 20px;
        padding: 1.5rem;
        border: 1px solid rgba(255,255,255,0.12);
        margin: 1rem 0;
    }

    /* 元素徽章 */
    .element-badge {
        display: inline-block;
        padding: 0.2rem 0.8rem;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
    }

    /* 清空聊天按钮 */
    .clear-chat-btn button {
        background: rgba(255,255,255,0.07) !important;
        border: 1px solid rgba(255,255,255,0.12) !important;
        border-radius: 12px !important;
        color: white !important;
        font-weight: 400 !important;
        box-shadow: none !important;
        height: 2.8rem !important;
        padding: 0 1rem !important;
    }
    .clear-chat-btn button:hover {
        background: rgba(255,255,255,0.15) !important;
        border-color: rgba(255,255,255,0.25) !important;
        transform: none !important;
        box-shadow: none !important;
    }

    /* 加载动画 */
    .stSpinner {
        width: 100% !important;
        text-align: center !important;
        padding: 1rem 0 !important;
    }
    .stSpinner > div {
        margin: 0 auto !important;
    }

    /* Toast 通知 */
    div[data-testid="stToast"] {
        color: black !important;
    }
    div[data-testid="stToast"] * {
        color: black !important;
    }
    </style>
    """, unsafe_allow_html=True)

# ============================================================
# 侧边栏组件
# ============================================================


def render_sidebar():
    """渲染侧边栏内容"""
    with st.sidebar:
        # 答题记录区域
        st.markdown(
            "<h2 style='color:white;font-size:1.3rem;'>📋 答题记录</h2>",
            unsafe_allow_html=True)

        if st.session_state.records is not None:
            st.dataframe(
                data=st.session_state.records,
                use_container_width=True,
                height=150)
            csv = st.session_state.records.to_csv(
                index=False, encoding='utf-8-sig')
            st.download_button(
                "📥 下载CSV",
                data=csv,
                file_name="元素人格测试_答题记录.csv",
                mime="text/csv",
                use_container_width=True
            )
        else:
            st.markdown(
                "<p style='color:rgba(255,255,255,0.5);font-size:0.85rem;'>完成测试后，答题记录将在此显示。</p>",
                unsafe_allow_html=True)

        # 历史记录（最近5次）
        if st.session_state.nickname:
            history = get_history(st.session_state.nickname, limit=5)
            if history:
                st.markdown(
                    "<hr style='margin:0.8rem 0;'>",
                    unsafe_allow_html=True)
                st.markdown(
                    "<p style='color:rgba(255,255,255,0.7);font-size:0.85rem;'>📈 历史测试（最近5次）</p>",
                    unsafe_allow_html=True)
                for idx, (ts, dom, w, f, wi, e) in enumerate(history):
                    try:
                        dt_obj = datetime.fromisoformat(
                            str(ts)) if ts else datetime.now()
                    except (ValueError, TypeError):
                        dt_obj = datetime.now()
                    st.markdown(
                        f"<p style='color:white;font-size:0.75rem;margin:0;'>#{
                            idx + 1} {
                            dt_obj.year}年{
                            dt_obj.month:02d}月{
                            dt_obj.day:02d}日 {
                            dt_obj.hour:02d}:{
                            dt_obj.minute:02d} · {dom}</p>",
                        unsafe_allow_html=True)

        # 导航按钮
        st.markdown("<hr style='margin:1rem 0;'>", unsafe_allow_html=True)

        if st.button("🔮 元素人格测试", use_container_width=True, key="sb_test"):
            st.session_state.page = "test"
            st.rerun()

        if st.button("📊 详细数据分析", use_container_width=True, key="sb_data"):
            st.session_state.prev_page = st.session_state.page
            st.session_state.page = "analysis"
            st.rerun()

        if st.button("🤖 元素人格助手", use_container_width=True, key="sb_chat"):
            st.session_state.prev_page = st.session_state.page
            st.session_state.page = "chat"
            st.rerun()

        if st.button("🌐 匿名元素社区", use_container_width=True, key="sb_community"):
            st.session_state.prev_page = st.session_state.page
            st.session_state.page = "community"
            st.rerun()

        if st.button(
            "💕 亲密关系占卜",
            use_container_width=True,
                key="sb_relationship"):
            st.session_state.prev_page = st.session_state.page
            st.session_state.page = "relationship"
            st.rerun()

        # 版本信息
        st.markdown("<hr style='margin:1rem 0;'>", unsafe_allow_html=True)
        st.markdown(
            "<p style='color:rgba(255,255,255,0.4);font-size:0.75rem;text-align:center;'>元素人格测试仪 v3.0</p>",
            unsafe_allow_html=True)


# ============================================================
# 随机昵称列表
# ============================================================
RANDOM_NAMES = [
    "元素旅人", "星辰探索者", "心灵拾荒者", "能量捕手", "平衡使者",
    "梦境漫游者", "灵魂画师", "宇宙尘埃", "光之追随者", "暗影舞者",
    "风之吟游诗人", "水之心", "火种少年", "大地守护者", "彩虹编织者",
    "月下独行者", "晨曦之子", "暮光旅人", "深海探索者", "云中漫步者"
]

# ============================================================
# 问卷组件
# ============================================================


def render_questionnaire():
    """渲染问卷（所有维度在同一页面）"""
    all_done = True
    total_q = len(QUESTIONS)
    done_q = sum(1 for q in QUESTIONS if st.session_state.get(
        f"q{q['id']:02d}") is not None)
    progress = done_q / total_q if total_q else 0

    st.progress(progress)
    st.markdown(
        f"<p style='text-align:center;color:rgba(255,255,255,0.5);font-size:0.85rem;'>{done_q}/{total_q} 已完成</p>",
        unsafe_allow_html=True)

    for elem in ELEMENT_ORDER:
        info = ELEMENTS[elem]
        qs = [q for q in QUESTIONS if q["element"] == elem]

        st.markdown(
            f"<div style='border-left:5px solid {info['color']};padding-left:16px;margin-top:1.2rem;'>"
            f"<h3 style='margin:0;padding:0;color:white;font-size:1.2rem;'>{info['emoji']} {elem}元素 · 人格维度</h3>"
            f"<p style='color:rgba(255,255,255,0.6);font-size:0.9rem;margin:4px 0 0 0;'>{info['desc']}</p>"
            f"</div>",
            unsafe_allow_html=True
        )
        for q in qs:
            key = f"q{q['id']:02d}"
            if q["type"] == "radio":
                st.radio(f"{q['id']}. {q['text']}", RADIO_OPTS,
                         key=key, index=None, horizontal=True)
            elif q["type"] == "select":
                st.selectbox(f"{q['id']}. {q['text']}", [
                             2, 4, 6, 8, 10], index=None, key=key)
            elif q["type"] == "slider":
                st.select_slider(f"{q['id']}. {q['text']}", options=[
                                 0, 2, 4, 6, 8, 10], value=None, key=key)
            elif q["type"] == "number":
                st.number_input(f"{q['id']}. {q['text']}",
                                0, 10, 0, 2, key=key)

    st.markdown("---")
    all_done = done_q == total_q
    if all_done:
        if st.button(
            "🌟 提交并揭示我的元素人格",
            use_container_width=True,
                type="primary"):
            vec, (scores_list, texts) = calc_scores()
            st.session_state.scores_raw = vec
            st.session_state._all_scores = scores_list
            st.session_state._question_texts = texts

            best, dists = compute_match(vec)
            st.session_state.best_match = best
            st.session_state.distances = dists
            st.session_state.scores = dict(zip(ELEMENT_ORDER, vec))

            st.session_state.records = pd.DataFrame({
                "题号": range(1, len(texts) + 1),
                "题目": texts,
                "分数": scores_list
            })

            nick = st.session_state.nickname or "匿名"
            save_record(nick, best, vec, json.dumps(scores_list))

            st.session_state.answer = None
            st.session_state.future = None

            st.session_state.page = "result"
            st.session_state.step = 0

            st.balloons()
            st.session_state.pending_toast = f"🎉 {nick}，问卷提交成功！正在生成你的元素人格报告 ✨"
            st.rerun()
    else:
        st.info(f"⚠️ 请完成所有 {total_q} 道题目后再提交")

# ============================================================
# 结果页面
# ============================================================


def render_results():
    """渲染测试结果页面"""
    # 显示待处理的提示消息
    if st.session_state.get("pending_toast"):
        st.toast(st.session_state.pending_toast, duration=3)
        st.session_state.pending_toast = ""

    # 获取测试结果数据
    element = st.session_state.best_match
    if not element:
        return

    vec = st.session_state.scores_raw
    short_elem = element.split()[-1] if " " in element else element

    # 获取元素信息
    emoji = next((v["emoji"]
                 for k, v in ELEMENTS.items() if k == short_elem), "✨")
    color = ELEMENTS[short_elem]["color"]
    desc = ELEMENTS[short_elem]["desc"]

    # 显示主导元素卡片
    st.markdown(f"""
    <div style='background:rgba(255,255,255,0.07);backdrop-filter:blur(20px);border-radius:24px;
                padding:2.5rem 2rem;text-align:center;border:1px solid rgba(255,255,255,0.12);margin:1.5rem 0;'>
        <div style='font-size:4.5rem;'>{emoji}</div>
        <div style='font-size:2rem;font-weight:800;color:{color};margin:0.5rem 0;'>{element}</div>
        <div style='color:white;font-size:0.95rem;line-height:1.8;max-width:500px;margin:0 auto;'>{desc}</div>
    </div>
    """, unsafe_allow_html=True)

    # 能量环形图
    fig = go.Figure(data=[go.Pie(
        labels=[f"{ELEMENTS[e]['emoji']} {e}" for e in ELEMENT_ORDER],
        values=[vec[i] for i in range(4)],
        hole=0.55,
        marker_colors=[ELEMENTS[e]["color"] for e in ELEMENT_ORDER],
        textinfo="label+percent",
        textfont=dict(color="white", size=13),
        hoverinfo="label+value+percent"
    )])
    fig.update_layout(
        paper_bgcolor="rgba(0,0,0,0)",
        plot_bgcolor="rgba(0,0,0,0)",
        font=dict(color="white"),
        height=350,
        margin=dict(l=10, r=10, t=10, b=10),
        showlegend=False
    )
    st.plotly_chart(fig, use_container_width=True, key="donut")

    # 元素能量分布四列显示（原始得分）
    st.markdown(
        "<h3 style='color:white;text-align:center;font-size:1.1rem;'>📊 元素能量分布</h3>",
        unsafe_allow_html=True)
    cols = st.columns(4)
    for col, elem in zip(cols, ELEMENT_ORDER):
        info = ELEMENTS[elem]
        val = vec[ELEMENT_ORDER.index(elem)]
        col.markdown(f"""
        <div style='text-align:center;padding:0.5rem;background:rgba(255,255,255,0.04);border-radius:12px;'>
            <div style='color:{info['color']};font-weight:700;font-size:1rem;'>{info['emoji']} {elem}</div>
            <div style='color:white;font-size:1.8rem;font-weight:800;'>{int(val)}</div>
            <div style='color:rgba(255,255,255,0.4);font-size:0.75rem;'>/ {info['max_score']}</div>
        </div>
        """, unsafe_allow_html=True)

    # 组合分析
    sorted_e = sorted(enumerate(vec), key=lambda x: -x[1])
    top2 = [ELEMENT_ORDER[i] for i in [sorted_e[0][0], sorted_e[1][0]]]
    compound = get_compound(*top2)

    if compound:
        st.markdown(f"""
        <div class='step-card'>
            <h4 style='color:white;'>🧬 元素组合分析</h4>
            <p style='color:white;font-size:1.1rem;font-weight:600;'>{compound['name']}</p>
            <p style='color:rgba(255,255,255,0.8);'>{compound['desc']}</p>
            <p style='color:rgba(255,255,255,0.7);font-size:0.9rem;'>{COMBO_SUG.get(compound['name'], '')}</p>
        </div>
        """, unsafe_allow_html=True)

    # 个性化建议
    sug = get_suggestions(vec)
    st.markdown(
        "<h3 style='color:white;text-align:center;'>💡 个性化成长建议</h3>",
        unsafe_allow_html=True)

    tab_labels = [f"{ELEMENTS[e]['emoji']} {e}" for e in ELEMENT_ORDER]
    try:
        default_idx = ELEMENT_ORDER.index(short_elem)
        tabs = st.tabs(tab_labels, default=default_idx)
    except (TypeError, ValueError):
        tabs = st.tabs(tab_labels)

    for i, elem in enumerate(ELEMENT_ORDER):
        with tabs[i]:
            for tip in sug[elem]:
                st.markdown(
                    f"<p style='color:white;text-align:center;'>• {tip}</p>",
                    unsafe_allow_html=True)

    # 能量小贴士
    tip = get_energy_tip(short_elem)
    st.markdown(f"""
    <div class='step-card' style='text-align:center;'>
        <p style='color:rgba(255,255,255,0.5);font-size:0.8rem;'>✨ 今日{emoji}元素能量小贴士</p>
        <p style='color:white;font-size:1rem;'>{tip}</p>
    </div>
    """, unsafe_allow_html=True)

    # 认知确认 + 第二人格探索
    st.session_state.answer = st.radio(
        "💭 这个结果符合你对自己的认知吗？",
        ["✓ 是的，非常符合", "✗ 不太符合，再看看"],
        horizontal=True,
        index=None
    )

    if st.session_state.answer == "✗ 不太符合，再看看":
        if st.button("🔮 探索我的潜在第二人格", use_container_width=True):
            # 按距离排序，取第二近的元素作为潜在第二人格
            sorted_d = sorted(
                st.session_state.distances.items(),
                key=lambda x: x[1])
            local_future = sorted_d[1][0] if len(
                sorted_d) >= 2 else sorted_d[0][0]
            st.session_state.future = local_future

            # 显示加载进度条
            bar = st.progress(0)
            for i in range(100):
                time.sleep(0.01)
                bar.progress(i + 1)

            st.toast("✨ 探索完成！")

        # 显示潜在第二人格结果
        if st.session_state.future:
            fc = ELEMENTS[st.session_state.future.split()[-1]]["color"]
            fd = ELEMENTS[st.session_state.future.split()[-1]]["desc"]
            fe = ELEMENTS[st.session_state.future.split()[-1]]["emoji"]
            st.markdown(f"""
            <div class='step-card' style='text-align:center;'>
                <div style='font-size:2.5rem;'>{fe}</div>
                <div style='font-size:1.3rem;font-weight:700;color:{fc};'>你的潜在第二人格：{st.session_state.future}</div>
                <div style='color:white;font-size:0.9rem;margin-top:0.5rem;'>{fd}</div>
            </div>
            """, unsafe_allow_html=True)

    elif st.session_state.answer == "✓ 是的，非常符合":
        st.success(f"🎉 太棒了！你已经找到了属于自己的元素力量！")

    # 操作按钮（下载报告、分享图片、匿名分享、重新测试）
    col1, col2, col3, col4 = st.columns([1, 1, 1, 1])

    with col1:
        nick = st.session_state.get("nickname", "") or "匿名用户"
        doc_bytes = generate_word(
            nick, element, vec, st.session_state.distances, sug)
        if st.download_button(
            "📄 下载报告",
            data=doc_bytes,
            file_name="元素人格报告.docx",
            key="dl_docx",
            use_container_width=True
        ):
            pass

    with col2:
        compound_name = compound['name'] if compound else None
        img_data = generate_result_image(nick, element, vec, compound_name)
        st.download_button(
            "🖼️ 分享图片",
            data=img_data,
            file_name=f"{nick}_元素人格.png",
            mime="image/png",
            use_container_width=True
        )

    with col3:
        if st.button(
            "🌐 匿名分享",
            use_container_width=True,
                key="share_community"):
            compound_name = compound['name'] if compound else None
            save_post(nick, element, vec, compound_name, "")
            st.session_state.pending_toast = f"📤 {nick}，已匿名分享到社区啦！去看看大家的元素之力吧 🌟"
            st.session_state.page = "community"
            st.rerun()

    with col4:
        if st.button("🔄 重新测试", use_container_width=True):
            # 清除测试相关状态
            for key in list(st.session_state.keys()):
                if key.startswith("q") or key in [
                    "best_match", "scores", "distances", "future", "records",
                    "answer", "scores_raw", "_all_scores", "_question_texts"
                ]:
                    del st.session_state[key]
            st.session_state.page = "test"
            st.rerun()

# ============================================================
# 聊天助手组件
# ============================================================


def render_chatbot():
    """渲染聊天助手页面"""
    # 标题栏（包含风格选择和清空按钮）
    col_title, col_style, col_clear = st.columns([3, 2, 1.2])

    with col_title:
        st.markdown(
            "<h2 style='color:white;margin:0;padding:0;'>🤖 元素人格助手</h2>",
            unsafe_allow_html=True)

    with col_style:
        style = st.selectbox(
            "",
            list(
                CHAT_STYLES.keys()),
            key="chat_style",
            label_visibility="collapsed")

    with col_clear:
        st.markdown("<div class='clear-chat-btn'>", unsafe_allow_html=True)
        if st.button("🗑️ 清空对话", key="clear_chat", use_container_width=True):
            st.session_state.chat_history = []
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

    # 动态快捷问题（仅在无对话时显示）
    if not st.session_state.chat_history:
        presets = dynamic_presets()
        st.markdown(
            "<p style='color:rgba(255,255,255,0.5);font-size:0.85rem;margin-bottom:0.3rem;'>💡 快捷提问</p>",
            unsafe_allow_html=True)
        pc = st.columns(3)
        for i, p in enumerate(presets):
            with pc[i]:
                if st.button(p, key=f"qpre_{i}", use_container_width=True):
                    # 添加用户消息
                    st.session_state.chat_history.append(
                        {"role": "user", "content": p})

                    # 构建系统提示词并调用 API
                    sys_prompt = build_system_prompt(style)
                    msgs = [{"role": "system", "content": sys_prompt}
                            ] + st.session_state.chat_history

                    with st.chat_message("assistant", avatar="🤖"):
                        with st.spinner("思考中..."):
                            reply = call_deepseek(msgs, False)
                        st.markdown(reply)

                    # 添加助手回复
                    st.session_state.chat_history.append(
                        {"role": "assistant", "content": reply})
                    st.rerun()

    # 聊天消息区域
    chat_box = st.container(height=400)
    with chat_box:
        if not st.session_state.chat_history:
            # 欢迎消息
            with st.chat_message("assistant", avatar="🤖"):
                st.markdown("你好！我是你的元素人格成长导师 🌟\n\n"
                            "可以尝试上方快捷问题，或直接告诉我你的困惑，"
                            "我会以当前风格与你交流探讨！")
        else:
            # 显示历史消息（最近30条）
            for msg in st.session_state.chat_history[-30:]:
                avatar = "🤖" if msg["role"] == "assistant" else "👤"
                with st.chat_message(msg["role"], avatar=avatar):
                    st.markdown(msg["content"])

    # 聊天输入框
    user_input = st.chat_input("输入你的问题...", key="main_chat")
    if user_input:
        # 添加用户消息
        st.session_state.chat_history.append(
            {"role": "user", "content": user_input})

        # 构建系统提示词并调用 API
        sys_prompt = build_system_prompt(style)
        msgs = [{"role": "system", "content": sys_prompt}] + \
            st.session_state.chat_history

        with st.chat_message("assistant", avatar="🤖"):
            with st.spinner("思考中..."):
                reply = call_deepseek(msgs, False)
            st.markdown(reply)

        # 添加助手回复
        st.session_state.chat_history.append(
            {"role": "assistant", "content": reply})

        st.rerun()

    st.markdown("---")

# ============================================================
# 数据分析页面
# ============================================================


def render_analysis():
    """渲染数据分析页面"""
    st.markdown(
        "<h2 style='color:white;margin-bottom:0.5rem;'>📊 详细数据分析</h2>",
        unsafe_allow_html=True)

    # 检查是否有测试数据
    if not st.session_state.distances or st.session_state.records is None or "scores_raw" not in st.session_state:
        st.info("请先完成测试并点击「🌟 揭示我的元素人格」，再查看数据分析。")
        st.markdown("---")
        return

    # 四个分析标签页
    tab1, tab2, tab3, tab4 = st.tabs(["📈 匹配度", "📋 每题得分", "🎯 能量雷达", "📉 历史趋势"])

    with tab1:
        st.markdown(
            "<h3 style='color:white;padding-top:0.8rem;'>四元素匹配度</h3>",
            unsafe_allow_html=True)
        total = sum(st.session_state.scores_raw)
        rows = []
        for name, dist in st.session_state.distances.items():
            elem_short = name.split()[-1]
            idx = ELEMENT_ORDER.index(
                elem_short) if elem_short in ELEMENT_ORDER else 0
            pct = round(
                st.session_state.scores_raw[idx] /
                total *
                100,
                1) if total else 0
            rows.append({"元素类型": name, "距离": f"{dist:.2f}", "匹配度": f"{pct}%"})
        df = pd.DataFrame(rows)
        st.dataframe(df, use_container_width=True, height=280)

    with tab2:
        # 每题得分趋势图
        st.markdown(
            "<h3 style='color:white;padding-top:0.8rem;'>每题得分趋势</h3>",
            unsafe_allow_html=True)
        recs = st.session_state.records

        fig = go.Figure()
        cmap = px.colors.sequential.Plasma_r
        marker_colors = px.colors.sample_colorscale(
            cmap, np.linspace(0, 1, 30))

        fig.add_trace(go.Scatter(
            x=recs["题号"],
            y=recs["分数"],
            mode="lines+markers",
            line=dict(color=cmap[3], width=2.5, shape="spline", smoothing=0.3),
            marker=dict(size=9, color=marker_colors,
                        line=dict(color="rgba(255,255,255,0.6)", width=1)),
            fill="tozeroy",
            fillcolor="rgba(156,39,176,0.08)",
            hovertemplate="题号 %{x}<br>分数: %{y}<extra></extra>"
        ))

        fig.update_layout(
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            font=dict(color="white"),
            height=350,
            xaxis=dict(title="题号", dtick=1, range=[0.5, 30.5], color="white"),
            yaxis=dict(title="分数", range=[-1, 11], color="white"),
            margin=dict(l=20, r=20, t=20, b=30)
        )

        st.plotly_chart(fig, use_container_width=True, key="trend")

        # 各元素得分汇总柱状图
        st.markdown(
            "<h3 style='color:white;'>各元素得分汇总</h3>",
            unsafe_allow_html=True)
        labs = [f"{ELEMENTS[e]['emoji']} {e}" for e in ELEMENT_ORDER]
        vals = [st.session_state.scores_raw[i] for i in range(4)]
        colors = [ELEMENTS[e]["color"] for e in ELEMENT_ORDER]

        fig2 = go.Figure(data=[go.Bar(
            x=labs,
            y=vals,
            marker_color=colors,
            text=[str(int(v)) for v in vals],
            textposition="outside",
            textfont=dict(color="white", size=13),
            hovertemplate="%{x}<br>得分: %{y}<extra></extra>"
        )])

        fig2.update_layout(
            paper_bgcolor="rgba(0,0,0,0)",
            plot_bgcolor="rgba(0,0,0,0)",
            font=dict(color="white"),
            height=350,
            yaxis=dict(title="得分", range=[0, max(vals) * 1.15], color="white"),
            margin=dict(l=20, r=20, t=20, b=30),
            showlegend=False
        )

        st.plotly_chart(fig2, use_container_width=True, key="bar_summary")

    with tab3:
        # 元素能量雷达图
        st.markdown(
            "<h3 style='color:white;padding-top:0.8rem;'>元素能量雷达图</h3>",
            unsafe_allow_html=True)
        cats = [f"{ELEMENTS[e]['emoji']} {e}" for e in ELEMENT_ORDER]
        vals = [min(int(st.session_state.scores_raw[i] /
                        ELEMENTS[ELEMENT_ORDER[i]]["max_score"] *
                        100), 100) for i in range(4)]

        fig3 = go.Figure(data=go.Scatterpolar(
            r=vals + [vals[0]],
            theta=cats + [cats[0]],
            fill="toself",
            fillcolor="rgba(102,126,234,0.25)",
            line=dict(color="#667eea", width=2.5),
            hovertemplate="%{theta}<br>能量: %{r}%<extra></extra>"
        ))

        fig3.update_layout(
            paper_bgcolor="rgba(0,0,0,0)",
            font=dict(
                color="white",
                size=13),
            polar=dict(
                radialaxis=dict(
                    range=[
                        0,
                        100],
                    visible=True,
                    color="rgba(255,255,255,0.3)"),
                bgcolor="rgba(0,0,0,0)",
                angularaxis=dict(
                    color="white",
                    tickfont=dict(
                        size=13))),
            height=450,
            margin=dict(
                l=40,
                r=40,
                t=20,
                b=20))

        st.plotly_chart(fig3, use_container_width=True, key="radar")

    with tab4:
        # 历史元素得分趋势
        st.markdown(
            "<h3 style='color:white;padding-top:0.8rem;'>历史元素得分趋势</h3>",
            unsafe_allow_html=True)

        if st.session_state.nickname:
            history = get_history(st.session_state.nickname, limit=10)

            if len(history) >= 2:
                df_hist = pd.DataFrame(
                    history, columns=[
                        "时间", "主导元素", "水", "火", "风", "土"])
                df_hist["时间"] = pd.to_datetime(
                    df_hist["时间"]).dt.strftime("%Y年%m月%d日 %H:%M")

                fig4 = go.Figure()
                for elem in ELEMENT_ORDER:
                    fig4.add_trace(
                        go.Scatter(
                            x=df_hist["时间"],
                            y=df_hist[elem],
                            mode='lines+markers',
                            name=elem,
                            line=dict(
                                color=ELEMENTS[elem]["color"],
                                width=3,
                                shape='spline',
                                smoothing=0.4),
                            marker=dict(
                                size=10,
                                color=ELEMENTS[elem]["color"],
                                line=dict(
                                    color='white',
                                    width=1.5)),
                            fill='tonexty',
                            fillcolor=f"rgba{
                                tuple(
                                    int(
                                        ELEMENTS[elem]['color'].lstrip('#')[
                                            j:j + 2],
                                        16) for j in (
                                        0,
                                        2,
                                        4)) + (
                                    0.15,
                                )}"))

                fig4.update_layout(
                    paper_bgcolor="rgba(0,0,0,0)",
                    plot_bgcolor="rgba(0,0,0,0)",
                    font=dict(
                        color="white"),
                    height=380,
                    xaxis=dict(
                        title="测试时间",
                        color="white",
                        showgrid=False),
                    yaxis=dict(
                        title="得分",
                        color="white",
                        showgrid=True,
                        gridcolor="rgba(255,255,255,0.08)"),
                    legend=dict(
                        font=dict(
                            color="white"),
                        orientation='h',
                        yanchor='bottom',
                        y=1.02,
                        xanchor='center',
                        x=0.5),
                    margin=dict(
                        l=20,
                        r=20,
                        t=50,
                        b=30),
                    hovermode='x unified')

                st.plotly_chart(
                    fig4,
                    use_container_width=True,
                    key="history_trend")
            else:
                st.info("至少需要2次测试记录才能显示趋势。")
        else:
            st.info("请设置昵称以查看历史趋势。")

    st.markdown("---")

# ============================================================
# 匿名社区组件
# ============================================================


def render_community():
    """渲染匿名社区页面"""
    # 显示待处理的提示消息
    if st.session_state.get("pending_toast"):
        st.toast(st.session_state.pending_toast, duration=3)
        st.session_state.pending_toast = ""

    # 页面标题
    st.markdown(
        "<h2 style='color:white;margin-bottom:0.5rem;'>🌐 匿名元素社区</h2>",
        unsafe_allow_html=True)
    st.markdown(
        "<p style='color:rgba(255,255,255,0.6);font-size:0.9rem;margin-bottom:1.5rem;'>在这里匿名分享你的元素人格，看看其他旅人的元素之力</p>",
        unsafe_allow_html=True)

    # 获取帖子列表并去重（每个昵称只保留最新一条）
    posts = get_posts(200)
    seen = set()
    unique_posts = []
    for p in posts:
        nick = p[1]
        if nick not in seen:
            seen.add(nick)
            unique_posts.append(p)

    if not unique_posts:
        st.info("还没有任何分享，快来发布第一条吧！")
        st.markdown("---")
        return

    # 渲染帖子列表
    for pid, nick, dom, w, f, wi, e, comp, msg, ts, likes in unique_posts[:50]:
        short_elem = dom.split()[-1] if " " in dom else dom
        info = ELEMENTS.get(short_elem, {})
        emoji = info.get("emoji", "✨")

        # 安全的日期解析
        try:
            dt_obj = datetime.fromisoformat(ts)
        except (ValueError, TypeError):
            dt_obj = datetime.now()

        ts_str = f"{
            dt_obj.year}年{
            dt_obj.month:02d}月{
            dt_obj.day:02d}日 {
                dt_obj.hour:02d}:{
                    dt_obj.minute:02d}"

        # 构建分数显示 HTML
        scores_html = "".join(
            f"<span style='margin:0 8px;'>{ELEMENTS[el]['emoji']} {el}: <strong>{int(v)}</strong></span>"
            for el, v in zip(ELEMENT_ORDER, [w, f, wi, e])
        )

        # 构建组合信息 HTML
        comp_html = f"<p style='color:rgba(255,255,255,0.5);font-size:0.85rem;margin:4px 0 0 0;'>🧬 {comp}</p>" if comp else ""

        # 渲染帖子卡片
        col_post, col_like = st.columns([10, 1])

        with col_post:
            st.markdown(f"""
            <div style='background:rgba(255,255,255,0.06);backdrop-filter:blur(10px);border-radius:16px;
                        padding:1.2rem 1.5rem;border:1px solid rgba(255,255,255,0.1);margin:0.8rem 0;'>
                <div style='display:flex;align-items:center;gap:10px;margin-bottom:6px;'>
                    <span style='font-size:1.8rem;'>{emoji}</span>
                    <span style='color:white;font-weight:600;font-size:1rem;'>{nick}</span>
                    <span style='color:rgba(255,255,255,0.35);font-size:0.8rem;margin-left:auto;'>{ts_str}</span>
                </div>
                <div style='color:white;font-size:0.9rem;'>{scores_html}</div>
                {comp_html}
            </div>
            """, unsafe_allow_html=True)

        with col_like:
            st.markdown(
                f"<div style='text-align:center;color:rgba(255,255,255,0.5);font-size:0.85rem;'>❤️ {likes}</div>",
                unsafe_allow_html=True)
            if pid not in st.session_state.liked_posts:
                if st.button("👍", key=f"like_{pid}", use_container_width=True):
                    like_post(pid)
                    st.session_state.liked_posts.add(pid)
                    st.rerun()
            else:
                st.markdown(
                    "<div style='text-align:center;color:#ff4d4d;font-size:1.2rem;'>❤️</div>",
                    unsafe_allow_html=True)

    st.markdown("---")

# ============================================================
# 亲密关系占卜
# ============================================================


def calc_scores_from_csv(uploaded_file):
    """从上传的CSV答题记录计算各元素得分"""
    df = pd.read_csv(uploaded_file)
    totals = {e: 0 for e in ELEMENT_ORDER}
    q_map = {q["id"]: q["element"] for q in QUESTIONS}
    for _, row in df.iterrows():
        qid = int(row["题号"])
        score = int(row["分数"])
        elem = q_map.get(qid)
        if elem:
            totals[elem] += score
    return [totals[e] for e in ELEMENT_ORDER]


RELATIONSHIP_ANALYSIS = {
    ("水", "水"): "双水组合 — 情感共鸣极强，彼此能深刻理解对方的情绪需求。缺点是容易陷入过度感性，需要有人适时拉回现实。",
    ("水", "火"): "水火相济 — 感性遇上行动力，互补性极强。水能软化火的急躁，火能带动水的激情，是经典的「吸引力组合」。",
    ("水", "风"): "流水遇风 — 交流畅快，精神层面高度契合。水的情感深度加上风的思维广度，能碰撞出大量创意火花。",
    ("水", "土"): "静水深流 — 水的包容遇上土的稳重，关系踏实而温暖。水需要土的可靠肩膀，土需要水的柔情滋养。",
    ("火", "火"): "双火组合 — 热情似火、能量爆棚，在一起永远不会无聊。但需注意控制脾气，避免「火上加火」的争吵。",
    ("火", "风"): "风助火势 — 火提供动力，风提供创意，是充满活力的组合。适合一起创业、旅行、探索未知。",
    ("火", "土"): "星火燎原 — 火的激情点燃土的潜力，土的稳重为火提供坚实后盾。是「梦想家+实干家」的黄金搭档。",
    ("风", "风"): "双风组合 — 思想自由奔放，在一起总有聊不完的话题。缺点是可能缺乏落地执行力，需要互相督促。",
    ("风", "土"): "风拂大地 — 风带来远方的视野，土提供脚下的根基。风帮土打破惯性，土帮风落实想法，非常平衡。",
    ("土", "土"): "双土组合 — 极度稳定可靠，彼此是最坚实的依靠。生活规律有序，但需注意偶尔制造惊喜，避免过于平淡。",
}


def calc_compatibility(vec_a, vec_b):
    """计算两个人的元素兼容性得分 (0-100)"""
    diffs = [abs(vec_a[i] - vec_b[i]) for i in range(4)]
    max_possible = 70
    avg_diff = sum(min(d, max_possible) for d in diffs) / 4
    score = max(0, 100 - avg_diff / max_possible * 100)
    return round(score, 1)


def get_relationship_report(elem_a, elem_b, vec_a, vec_b):
    """生成亲密关系分析报告"""
    key = tuple(sorted([elem_a, elem_b]))
    base_analysis = RELATIONSHIP_ANALYSIS.get(key, "元素组合分析待补充。")

    dominant_a = max(range(4), key=lambda i: vec_a[i])
    dominant_b = max(range(4), key=lambda i: vec_b[i])

    insights = []
    if vec_a[dominant_a] > vec_b[dominant_b] * 1.5:
        insights.append(
            f"在{ELEMENT_ORDER[dominant_a]}元素维度上，双方的差异较大，建议多理解对方的表达方式。")
    if abs(vec_a[1] - vec_b[1]) < 10:
        insights.append("在火元素维度上能量相近，意味着双方的行动力和热情水平匹配良好。")
    total = sum(vec_a) + sum(vec_b)
    if total > 200:
        insights.append("两人的元素总能量充沛，关系充满活力与可能性！")
    elif total < 100:
        insights.append("两人的元素能量偏内敛，关系安静而深沉，需要主动营造交流氛围。")

    compat_score = calc_compatibility(vec_a, vec_b)
    return base_analysis, insights, compat_score


def render_relationship():
    """渲染亲密关系占卜页面"""
    st.markdown(
        "<h2 style='color:white;margin-bottom:0.5rem;'>💕 亲密关系占卜</h2>",
        unsafe_allow_html=True)
    st.markdown(
        "<p style='color:rgba(255,255,255,0.6);font-size:0.9rem;margin-bottom:1.5rem;'>"
        "上传两份答题记录（测试对象 & 你自己），探索你们的元素兼容性</p>",
        unsafe_allow_html=True)

    col1, col2 = st.columns(2)
    with col1:
        file_a = st.file_uploader("👤 测试对象的答题记录", type="csv", key="rel_a")
    with col2:
        file_b = st.file_uploader("👤 你自己的答题记录", type="csv", key="rel_b")

    if file_a and file_b:
        try:
            vec_a = calc_scores_from_csv(file_a)
            vec_b = calc_scores_from_csv(file_b)
        except Exception as e:
            st.error(f"CSV 解析失败，请确认文件格式正确：{e}")
            return

        # 计算主导元素
        elem_a = ELEMENT_ORDER[max(range(4), key=lambda i: vec_a[i])]
        elem_b = ELEMENT_ORDER[max(range(4), key=lambda i: vec_b[i])]

        st.markdown("---")
        st.markdown(
            "<h3 style='color:white;text-align:center;'>📊 元素能量对比</h3>",
            unsafe_allow_html=True)

        # 并排展示两人元素得分
        c1, c2 = st.columns(2)
        for col, label, vec in [(c1, "测试对象", vec_a), (c2, "你自己", vec_b)]:
            col.markdown(
                f"<p style='color:white;text-align:center;font-weight:600;'>{label}</p>",
                unsafe_allow_html=True)
            fig = go.Figure(data=[go.Pie(
                labels=[f"{ELEMENTS[e]['emoji']} {e}" for e in ELEMENT_ORDER],
                values=vec,
                hole=0.55,
                marker_colors=[ELEMENTS[e]["color"] for e in ELEMENT_ORDER],
                textinfo="label+percent",
                textfont=dict(color="white", size=11),
            )])
            fig.update_layout(
                paper_bgcolor="rgba(0,0,0,0)", plot_bgcolor="rgba(0,0,0,0)", font=dict(
                    color="white"), height=250, margin=dict(
                    l=10, r=10, t=10, b=10), showlegend=False)
            col.plotly_chart(
                fig,
                use_container_width=True,
                key=f"rel_donut_{label}")

        # 兼容性分数
        compat = calc_compatibility(vec_a, vec_b)
        compat_color = "#4fc3f7" if compat >= 70 else "#ffd54f" if compat >= 40 else "#ff7043"
        st.markdown(f"""
        <div style='text-align:center;padding:1.5rem;background:rgba(255,255,255,0.06);border-radius:20px;margin:1rem 0;'>
            <p style='color:rgba(255,255,255,0.5);font-size:0.85rem;'>💕 元素兼容度评分</p>
            <p style='color:{compat_color};font-size:3.5rem;font-weight:800;'>{compat}</p>
            <p style='color:rgba(255,255,255,0.7);font-size:0.95rem;'>
                {"✨ 高度契合！" if compat >= 70 else "🌟 有一定默契" if compat >= 40 else "💪 需要更多磨合"}
            </p>
        </div>
        """, unsafe_allow_html=True)

        # 主导元素组合分析
        report_text, insights, _ = get_relationship_report(
            elem_a, elem_b, vec_a, vec_b)
        st.markdown(f"""
        <div style='background:rgba(255,255,255,0.06);border-radius:20px;padding:1.5rem;margin:1rem 0;
                    border-left:5px solid {ELEMENTS[elem_a]['color']};'>
            <p style='color:white;font-size:1rem;font-weight:600;'>
                {ELEMENTS[elem_a]['emoji']} {elem_a}  +  {ELEMENTS[elem_b]['emoji']} {elem_b}
            </p>
            <p style='color:rgba(255,255,255,0.8);font-size:0.95rem;line-height:1.7;'>{report_text}</p>
        </div>
        """, unsafe_allow_html=True)

        # 深入洞察
        if insights:
            st.markdown(
                "<h4 style='color:white;'>🔍 深入洞察</h4>",
                unsafe_allow_html=True)
            for tip in insights:
                st.markdown(
                    f"<p style='color:rgba(255,255,255,0.8);font-size:0.9rem;'>• {tip}</p>",
                    unsafe_allow_html=True)

        # 元素细节对比表
        st.markdown(
            "<h4 style='color:white;margin-top:1rem;'>📋 元素细节对比</h4>",
            unsafe_allow_html=True)
        comp_rows = []
        for i, elem in enumerate(ELEMENT_ORDER):
            info = ELEMENTS[elem]
            comp_rows.append({
                "元素": f"{info['emoji']} {elem}",
                "测试对象得分": vec_a[i],
                "你的得分": vec_b[i],
                "差值": abs(vec_a[i] - vec_b[i])
            })
        st.dataframe(
            pd.DataFrame(comp_rows),
            use_container_width=True,
            height=200)

    elif file_a or file_b:
        st.info("请上传两份答题记录以进行对比分析")

    st.markdown("---")

# ============================================================
# 主流程
# ============================================================


def main():
    """主入口函数"""
    render_css()
    render_sidebar()

    # 标题区域
    st.markdown(
        "<div style='text-align:center;padding:1.5rem 0 0.5rem 0;'>"
        "<h1>🔮 元素人格测试仪</h1>"
        "<p style='color:rgba(255,255,255,0.8);font-size:1.05rem;letter-spacing:2px;'>"
        "发现你内在的四大元素之力 · 水 💧 火 🔥 风 🌬️ 土 🌍</p></div>",
        unsafe_allow_html=True)

    st.markdown("---")

    page = st.session_state.page

    # 测试/结果页面
    if page in ("test", "result"):
        if "scores_raw" not in st.session_state:
            # 初始化昵称
            if "nickname" not in st.session_state or not st.session_state.nickname:
                st.session_state.nickname = random.choice(RANDOM_NAMES)
            if "nick_rk" not in st.session_state:
                st.session_state.nick_rk = 0

            # 昵称输入区域
            nc1, nc2 = st.columns([5, 1])
            with nc1:
                nick = st.text_input(
                    "💫 你的昵称：",
                    value=st.session_state.nickname,
                    key=f"nick_{
                        st.session_state.nick_rk}")
                st.session_state.nickname = nick
            with nc2:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button(
                    "🎲 换一个",
                    use_container_width=True,
                        key="nick_btn"):
                    st.session_state.nick_rk += 1
                    st.session_state.nickname = random.choice(RANDOM_NAMES)
                    st.rerun()

            # 同意条款
            if st.checkbox("✓ 测试结果仅供娱乐参考，我已知情", key="agree_check"):
                st.success(
                    f"✨ {
                        st.session_state.nickname}，准备好探索你的元素人格了吗？开始答题吧！")

            st.markdown("---")
            render_questionnaire()
        else:
            render_results()

    # 聊天助手页面
    elif page == "chat":
        col_back, _ = st.columns([1, 8])
        with col_back:
            if st.button("⬅ 返回", use_container_width=True):
                st.session_state.page = st.session_state.prev_page
                st.rerun()
        render_chatbot()

    # 匿名社区页面
    elif page == "community":
        col_back, _ = st.columns([1, 8])
        with col_back:
            if st.button("⬅ 返回", use_container_width=True):
                st.session_state.page = st.session_state.prev_page
                st.rerun()
        render_community()

    # 数据分析页面
    elif page == "analysis":
        col_back, _ = st.columns([1, 8])
        with col_back:
            if st.button("⬅ 返回", use_container_width=True):
                st.session_state.page = st.session_state.prev_page
                st.rerun()
        render_analysis()

    # 亲密关系占卜页面
    elif page == "relationship":
        col_back, _ = st.columns([1, 8])
        with col_back:
            if st.button("⬅ 返回", use_container_width=True):
                st.session_state.page = st.session_state.prev_page
                st.rerun()
        render_relationship()

    # 页脚
    st.markdown(
        "<div style='text-align:center;padding:1.5rem 0;color:rgba(255,255,255,0.4);font-size:0.8rem;'>"
        "🔮 元素人格测试仪 · 探索内在的四大元素之力 · v3.0</div>",
        unsafe_allow_html=True)


if __name__ == "__main__":
    main()
